from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

class ResultPage:
    def __init__(self, parent, app, time_complete=0, correct_answers=0, total_questions=0, user_name=None,
                 test_name=None, user_answers=None):
        self.parent = parent
        self.app = app

        self.user_answers = user_answers
        self.time_complete = time_complete
        self.total_questions = total_questions
        self.correct_answers = correct_answers
        self.user_name = user_name
        self.correct_percent = correct_answers / total_questions * 100
        self.certificates_dir = "certificates"

        print(f"ResultPage инициализирован с user_name: {self.user_name}")

        # Берём данные теста, если есть имя
        self.test_name = test_name
        self.test_data = self.app.test_manager.get_tests(test_name) if test_name else None
        self.grade = None

        self.create_widgets()
        self.label_grade_change()
        self.save_results_to_db()
        self.show_detail_info()

    def create_widgets(self):
        # Надпись "Результаты тестирования"
        self.label_title = ttk.Label(self.parent,
                                     text="Результаты тестирования",
                                     font=('Arial', 17, 'bold'),
                                     style='Label.TLabel')
        self.label_title.pack(anchor='w', padx=20, pady=10)

        # Фрейм для всего
        self.body_frame = Frame(self.parent,
                                relief=SOLID,
                                borderwidth=2,
                                background='white')
        self.body_frame.pack(fill=BOTH, expand=True, padx=20, pady=20)

        # Кнопка "Назад в меню"
        self.menu_button = ttk.Button(self.parent,
                                      text='Назад в меню',
                                      style='StyleGray.TButton',
                                      command=self.back_to_menu)
        self.menu_button.pack(side=LEFT, padx=20, pady=10)

        # Кнопка для скачивания сертификата в word
        self.DownloadWordBtn = ttk.Button(self.parent,
                                      text='Скачать сертификат word',
                                      style='StyleBlue.TButton',
                                      command=self.download_certificate_word)
        self.DownloadWordBtn.pack(side=RIGHT, padx=20, pady=10)

        # Кнопка для скачивания сертификата в pdf
        self.DownloadPDFBtn = ttk.Button(self.parent,
                                         text='Скачать сертификат pdf',
                                         style='StyleRed.TButton',
                                         command=self.download_certificate_pdf)
        self.DownloadPDFBtn.pack(side=RIGHT, pady=10)

        # Оценка
        self.label_grade = ttk.Label(self.body_frame,
                                     text="Результат: ",
                                     font=('Arial', 15, 'bold'),
                                     style='Label.TLabel')
        self.label_grade.pack(pady=10)

        # Количество правильных вариантов ответа
        self.label_true_answers = ttk.Label(self.body_frame,
                                            text=f"Правильных ответов: {self.correct_answers} из {self.total_questions} ({self.correct_percent}%)",
                                            font=('Arial', 15, 'bold'),
                                            style='Label.TLabel')
        self.label_true_answers.pack(pady=3)

        # Время прохождения
        self.label_travel_time = ttk.Label(self.body_frame,
                                           text=f"Время прохождения: {self.format_time()}",
                                           font=('Arial', 14,),
                                           style='Label.TLabel')
        self.label_travel_time.pack(pady=3)

        # Надпись "Детальный результат"
        self.label_detail = ttk.Label(self.body_frame,
                                      text="Детальный результат: ",
                                      font=('Arial', 15, 'bold'),
                                      style='Label.TLabel')
        self.label_detail.pack(anchor='w', padx=10)

        # Фрейм для детального результата
        self.detail_frame = Frame(self.body_frame,
                                  relief=SOLID,
                                  borderwidth=1,
                                  background='white')
        self.detail_frame.pack(expand=True, fill=BOTH, padx=10, pady=10)

        # Canvas and scrollbar для прокрутки
        self.canvas = Canvas(self.detail_frame,
                             background="white",
                             highlightthickness=0)
        self.scrollbar = Scrollbar(self.detail_frame,
                                   orient='vertical',
                                   command=self.canvas.yview)
        self.scrollable_frame = Frame(self.canvas,
                                      background='white')

        self.scrollable_frame.bind('<Configure>',
                                   lambda e: self.canvas.configure(scrollregion=self.canvas.bbox('all')))
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor='nw')
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        # Упаковываем scrollbar и canvas
        self.scrollbar.pack(side=RIGHT, fill=Y)
        self.canvas.pack(side=LEFT, fill=BOTH, expand=True)

    def show_detail_info(self):
        if not self.test_data or not self.user_answers:
            no_data_label = ttk.Label(self.scrollable_frame,
                                      text="Детальная информация недоступна",
                                      font=('Arial', 12),
                                      background='white')
            no_data_label.pack(pady=20)
            return

        # Устанавливаем минимальную ширину для scrollable_frame
        self.scrollable_frame.update_idletasks()  # Обновляем геометрию
        canvas_width = self.canvas.winfo_width()
        if canvas_width < 100:  # Если canvas еще не отрисовался
            canvas_width = 800  # Устанавливаем разумную ширину по умолчанию

        # Настраиваем scrollable_frame на всю ширину canvas
        self.canvas.itemconfig(self.canvas.find_all()[0], width=canvas_width)

        for count_questions in range(self.total_questions):
            questions_key = f"question_{count_questions + 1}"
            if questions_key in self.test_data:
                questions_data = self.test_data[questions_key]
                user_answer_index = self.user_answers[count_questions] if count_questions < len(
                    self.user_answers) else -1
                correct_answer_index = questions_data["correct_answer"]

                # Определяем, правильный ли ответ
                is_correct = user_answer_index == correct_answer_index

                # Создание фрейма для вопроса
                question_frame = Frame(self.scrollable_frame,
                                       background='#D7D7D7',
                                       relief=SOLID,
                                       borderwidth=1,
                                       width=800)
                question_frame.pack(fill=X, expand=True, padx=5, pady=5, ipady=10)
                question_frame.bind('<MouseWheel>', self._on_mousewheel)

                # Заголовок вопроса
                question_label = ttk.Label(question_frame,
                                           text=f"Вопрос {count_questions + 1}: {questions_data['question']}",
                                           font=('Arial', 14, 'bold'),
                                           background='#D7D7D7',
                                           wraplength=800,
                                           justify=LEFT)
                question_label.pack(anchor='w', padx=10, pady=10)
                question_label.bind('<MouseWheel>', self._on_mousewheel)

                # Правильный ответ
                correct_answer_text = questions_data['options'][correct_answer_index]

                correct_label = ttk.Label(question_frame,
                                          text=f'Правильный ответ: {correct_answer_text}',
                                          font=('Cabin', 12),
                                          background="#D7D7D7",
                                          foreground="#444444",
                                          wraplength=800,
                                          justify=LEFT)
                correct_label.pack(anchor='w', padx=10)
                correct_label.bind('<MouseWheel>', self._on_mousewheel)

                # Ответ пользователя
                if user_answer_index != -1:
                    user_answer_text = questions_data['options'][user_answer_index]

                    user_label = Label(question_frame,
                                       text=f"Ваш ответ: {user_answer_text}",
                                       font=('Cabin', 12),
                                       background='#D7D7D7',
                                       foreground="#444444",
                                       justify=LEFT)
                    user_label.pack(anchor='w', padx=10)
                    user_label.bind('<MouseWheel>', self._on_mousewheel)

                else:
                    user_label = Label(question_frame,
                                       text="Ваш ответ: Не отвечено",
                                       font=('Arial', 11),
                                       background='#D7D7D7',
                                       foreground='red',
                                       justify=LEFT)
                    user_label.pack(anchor='w', padx=10, pady=2)
                    user_label.bind('<MouseWheel>', self._on_mousewheel)

                # Статус Правильно\Неправильно
                status_text = "✓ Правильно" if is_correct else "✗ Неправильно"
                status_color = 'green' if is_correct else 'red'

                status_label = ttk.Label(question_frame,
                                         text=status_text,
                                         foreground=status_color,
                                         font=('Arial', 13, 'bold'),
                                         background='#D7D7D7')
                status_label.pack(anchor='w', padx=10, pady=2)

        # Привязываем прокрутку к основным контейнерам
        self.scrollable_frame.bind('<MouseWheel>', self._on_mousewheel)
        self.canvas.bind('<MouseWheel>', self._on_mousewheel)

        # Обновляем область прокрутки
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

    def back_to_menu(self):
        from pages.main_menu import MainMenu
        self.app.show_page(MainMenu)

    def _on_mousewheel(self, event: Event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')

    def get_grade(self):
        if self.correct_percent <= 30:
            self.grade = "Плохо"

        elif self.correct_percent < 80:
            self.grade = "Хорошо"

        else:
            self.grade = "Отлично"

    def label_grade_change(self):
        self.get_grade()
        if self.grade == "Плохо":
            self.label_grade.config(text=f"Результат: {self.grade}",
                                    foreground='red')
        elif self.grade == "Хорошо":
            self.label_grade.config(text=f"Результат: {self.grade}",
                                    foreground='#f0bc00')
        else:
            self.label_grade.config(text=f"Результат: {self.grade}",
                                    foreground='green')

    def format_time(self):
        self.minutes, self.secs = divmod(self.time_complete, 60)
        return f'{self.minutes:02d}:{self.secs:02d}'

    def save_results_to_db(self):
        import requests

        # Данные для отправки
        self.result_data = {
            "user_name": self.user_name,
            "test_name": self.test_name,
            "total_questions": self.total_questions,
            "correct_answers": self.correct_answers,
            "percent_correct_answers": round(self.correct_percent, 2),
            "time_complete": self.time_complete,
            "user_answers": self.user_answers
        }

        try:
            # URL FastAPI endpoint
            api_url = 'http://127.0.0.1:8000/api/save_test_result'

            response = requests.post(
                api_url,
                json=self.result_data,
                headers={'Content-Type': 'application/json'},
                timeout=10)

            # Проверка статуса ответа
            if response.status_code == 200:
                result = response.json()
                print("Результаты успешно сохранены через FastAPI")
                print(f"ID результата: {result.get('result_id')}")
                print(f"ID пользователя: {result.get('user_id')}")
            else:
                error_detail = response.json().get('detail', 'Unknown error')
                print(f"Ошибка при сохранении результатов: {response.status_code}")
                print(f"Детали ошибки: {error_detail}")

        except requests.exceptions.ConnectionError:
            print("Ошибка подключения: Не удалось соединиться с FastAPI сервером")
            print("Убедитесь, что сервер FastAPI запущен на localhost:8000")
        except requests.exceptions.Timeout:
            print("Ошибка подключения: Превышено время ожидания сервера")
        except requests.exceptions.RequestException as e:
            print(f"Ошибка при отправке запроса: {e}")
        except Exception as e:
            print(f"Неожиданная ошибка: {e}")

    def download_certificate_word(self):
        try:
            template_path = "certificates/sample.docx"

            doc = Document(template_path)

            certificate_number = datetime.now().strftime("%Y%m%d%H%M")

            replacements = {
                "{{CERTIFICATE_NUMBER}}": certificate_number,
                "{{USER_NAME}}": self.user_name or "Пользователь",
                "{{TEST_NAME}}": self.test_name or "Тест",
                "{{CORRECT_ANSWERS}}": str(self.correct_answers),
                "{{TOTAL_QUESTIONS}}": str(self.total_questions),
                "{{PERCENTAGE}}": f"{str(self.correct_percent)}%",
                "{{GRADE}}": self.grade,
                "{{TIME}}": str(self.time_complete),
                "{{DATE}}": datetime.now().strftime('%d.%m.%Y %H:%M')
            }

            # Замена текста
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    for run in paragraph.runs:
                                        if key in run.text:
                                            run.text = run.text.replace(key, value)

            # Сохранение результатов
            filename = f"certificate_{self.test_name}_{certificate_number}.docx"
            filepath = os.path.join(self.certificates_dir, filename)

            doc.save(filepath)

            messagebox.showinfo("Успех", f"Сертификат успешно сохранён:\n{filename}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить сертификат:\n{e}")

    def download_certificate_pdf(self):
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"certificate_{self.user_name}_{timestamp}.pdf"
            filepath = os.path.join(self.certificates_dir, filename)

            # Стандартные пути к шрифтам в Windows
            windows_fonts_paths = [
                "C:/Windows/Fonts/arial.ttf",
                "C:/Windows/Fonts/ARIAL.TTF",
                "C:/Windows/Fonts/ariali.ttf",  # Arial Italic
                "C:/WinNT/Fonts/arial.ttf",  # Для старых версий Windows
            ]

            windows_bold_paths = [
                "C:/Windows/Fonts/arialbd.ttf",
                "C:/Windows/Fonts/ARIALBD.TTF",
                "C:/Windows/Fonts/arialbi.ttf",  # Arial Bold Italic
                "C:/WinNT/Fonts/arialbd.ttf",
            ]

            # Поиск обычного шрифта Arial
            font_path = None
            for path in windows_fonts_paths:
                if os.path.exists(path):
                    font_path = path
                    break

            # Поиск жирного шрифта Arial
            font_bold_path = None
            for path in windows_bold_paths:
                if os.path.exists(path):
                    font_bold_path = path
                    break

            # Регистрируем найденные шрифты
            font_name = 'Helvetica'  # значение по умолчанию
            font_bold = 'Helvetica-Bold'

            if font_path and font_bold_path:
                try:
                    pdfmetrics.registerFont(TTFont('Arial', font_path))
                    pdfmetrics.registerFont(TTFont('Arial-Bold', font_bold_path))
                    font_name = 'Arial'
                    font_bold = 'Arial-Bold'
                    print(f"Используются шрифты: {font_path}, {font_bold_path}")
                except Exception as e:
                    print(f"Ошибка регистрации шрифтов: {e}")
                    # Пробуем зарегистрировать только обычный шрифт
                    try:
                        pdfmetrics.registerFont(TTFont('Arial', font_path))
                        font_name = 'Arial'
                        font_bold = 'Helvetica-Bold'
                    except:
                        pass
            else:
                print("Шрифты Arial не найдены, используем стандартные Helvetica")

            # Альтернативные варианты если Arial не найден
            if font_name == 'Helvetica':
                # Пробуем найти другие кириллические шрифты
                alternative_fonts = [
                    "C:/Windows/Fonts/tahoma.ttf",
                    "C:/Windows/Fonts/tahomabd.ttf",
                    "C:/Windows/Fonts/verdana.ttf",
                    "C:/Windows/Fonts/verdanab.ttf",
                    "C:/Windows/Fonts/times.ttf",
                    "C:/Windows/Fonts/timesbd.ttf",
                ]

                for i in range(0, len(alternative_fonts), 2):
                    if i + 1 < len(alternative_fonts) and os.path.exists(alternative_fonts[i]):
                        try:
                            regular_path = alternative_fonts[i]
                            bold_path = alternative_fonts[i + 1] if os.path.exists(alternative_fonts[i + 1]) else \
                            alternative_fonts[i]

                            font_family = os.path.basename(regular_path).split('.')[0].capitalize()
                            pdfmetrics.registerFont(TTFont(font_family, regular_path))
                            pdfmetrics.registerFont(TTFont(font_family + '-Bold', bold_path))

                            font_name = font_family
                            font_bold = font_family + '-Bold'
                            print(f"Используются альтернативные шрифты: {regular_path}, {bold_path}")
                            break
                        except Exception as e:
                            print(f"Ошибка регистрации альтернативного шрифта: {e}")
                            continue

            c = canvas.Canvas(filepath, pagesize=A4)
            width, height = A4

            # Фон
            c.setFillColorRGB(0.98, 0.98, 0.98)
            c.rect(0, 0, width, height, fill=1)
            c.setFillColorRGB(0, 0, 0)

            # Рамка
            c.setStrokeColorRGB(0.3, 0.5, 0.8)
            c.setLineWidth(2)
            c.rect(1.5 * cm, 1.5 * cm, width - 3 * cm, height - 3 * cm)

            # Заголовок
            c.setFont(font_bold, 26)
            c.setFillColorRGB(0.2, 0.4, 0.8)
            c.drawCentredString(width / 2, height - 4 * cm, "СЕРТИФИКАТ")

            # Основной текст
            c.setFont(font_name, 14)
            c.setFillColorRGB(0, 0, 0)
            c.drawCentredString(width / 2, height - 7 * cm, "Настоящим удостоверяется, что")

            # Имя пользователя
            c.setFont(font_bold, 16)
            c.setFillColorRGB(0.8, 0.2, 0.2)
            c.drawCentredString(width / 2, height - 9 * cm, self.user_name or "Пользователь")

            c.setFont(font_name, 14)
            c.setFillColorRGB(0, 0, 0)
            c.drawCentredString(width / 2, height - 11 * cm, "успешно прошел(ла) тестирование")
            c.drawCentredString(width / 2, height - 12.5 * cm, "по курсу:")

            # Название теста
            c.setFont(font_bold, 15)
            c.setFillColorRGB(0.2, 0.4, 0.8)
            test_name_text = f"«{self.test_name}»" if self.test_name else "«Неизвестный тест»"
            c.drawCentredString(width / 2, height - 14.5 * cm, test_name_text)

            # Результаты
            results_y = height - 18 * cm

            c.setFont(font_bold, 14)
            c.drawCentredString(width / 2, results_y, "Результаты тестирования")

            c.setFont(font_name, 12)
            c.drawString(4 * cm, results_y - 1.5 * cm,
                         f"Правильных ответов: {self.correct_answers} из {self.total_questions}")
            c.drawString(4 * cm, results_y - 2.5 * cm, f"Процент выполнения: {self.correct_percent:.1f}%")
            c.drawString(4 * cm, results_y - 3.5 * cm, f"Оценка: {self.grade}")
            c.drawString(4 * cm, results_y - 4.5 * cm, f"Время выполнения: {self.format_time()}")

            # Дата
            c.drawString(4 * cm, results_y - 6 * cm, f"Дата прохождения: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

            # Подпись
            signature_y = 6 * cm
            c.setFont(font_name, 12)
            c.drawString(width - 8 * cm, signature_y + 1.5 * cm, "Директор по обучению")
            c.line(width - 8 * cm, signature_y + 1 * cm, width - 3 * cm, signature_y + 1 * cm)

            c.save()

            messagebox.showinfo("Успех", f"PDF сертификат сохранен:\n{filename}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать PDF сертификат:\n{str(e)}")

    # Очистка виджетов при закрытии страницы
    def destroy(self):
        for widget in self.parent.winfo_children():
            widget.destroy()
